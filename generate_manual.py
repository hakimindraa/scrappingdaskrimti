"""
Script untuk generate Manual Book DASTA (Deskripsi Data) dalam format Word (.docx)
Fokus: Tutorial Penggunaan Aplikasi Website — Bab per Bab
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "064E3B")
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0FDF4")
    return table


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"💡 Tips: {text}")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(4, 78, 59)


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ Perhatian: {text}")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(180, 80, 0)


def add_important(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"❗ Penting: {text}")
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)


def add_numbered_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')


def create_manual():
    doc = Document()

    # ===================== STYLES =====================
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(4, 78, 59)

    # ===================== COVER PAGE =====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL BOOK")
    run.bold = True
    run.font.size = Pt(38)
    run.font.color.rgb = RGBColor(4, 78, 59)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("DASTA — Deskripsi Data")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(6, 95, 70)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Panduan Penggunaan Aplikasi Website")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(3):
        doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Versi Dokumen", "2.0"),
        ("Tanggal", datetime.date.today().strftime("%d %B %Y")),
        ("Klasifikasi", "Internal"),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.rows[i].cells[0]
        cell_r = info_table.rows[i].cells[1]
        cell_l.text = label
        cell_r.text = value
        for p in cell_l.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        for p in cell_r.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
        set_cell_shading(cell_l, "D1FAE5")

    doc.add_page_break()

    # ===================== DAFTAR ISI =====================
    doc.add_heading('DAFTAR ISI', level=1)
    toc_items = [
        ("BAB 1", "Mengenal Aplikasi DASTA", ""),
        ("  1.1", "Apa itu DASTA?", ""),
        ("  1.2", "Fitur Utama", ""),
        ("  1.3", "Struktur Menu Navigasi", ""),
        ("BAB 2", "Halaman Utama (Home)", ""),
        ("  2.1", "Tampilan Selamat Datang", ""),
        ("  2.2", "Kartu Status Ringkasan", ""),
        ("  2.3", "Kartu Status Per Scraper", ""),
        ("  2.4", "Status Server & Response Time", ""),
        ("  2.5", "Activity Log", ""),
        ("  2.6", "Fitur Auto-Refresh (LIVE)", ""),
        ("BAB 3", "Scraping SIPEDE", ""),
        ("  3.1", "Persiapan Sebelum Scraping", ""),
        ("  3.2", "Langkah 1 — Buka Browser", ""),
        ("  3.3", "Langkah 2 — Login ke SIPEDE", ""),
        ("  3.4", "Langkah 3 — Konfigurasi & Mulai Scraping", ""),
        ("  3.5", "Langkah 4 — Melihat Progress Scraping", ""),
        ("  3.6", "Langkah 5 — Hasil Scraping & Export Data", ""),
        ("  3.7", "Scraping Tahun Lain", ""),
        ("BAB 4", "Scraping SPDP", ""),
        ("  4.1", "Persiapan Sebelum Scraping", ""),
        ("  4.2", "Langkah 1 — Buka Browser", ""),
        ("  4.3", "Langkah 2 — Login & Navigasi ke Halaman Data", ""),
        ("  4.4", "Langkah 3 — Deteksi Tabel & Konfigurasi", ""),
        ("  4.5", "Langkah 4 — Melihat Progress Scraping", ""),
        ("  4.6", "Langkah 5 — Hasil Scraping & Export Data", ""),
        ("BAB 5", "Scraping DASTI", ""),
        ("  5.1", "Persiapan Sebelum Scraping", ""),
        ("  5.2", "Langkah 1 — Buka Browser", ""),
        ("  5.3", "Langkah 2 — Login ke DASTI (dengan Captcha)", ""),
        ("  5.4", "Langkah 3 — Navigasi ke Halaman Data", ""),
        ("  5.5", "Langkah 4 — Konfigurasi & Mulai Scraping", ""),
        ("  5.6", "Langkah 5 — Hasil Scraping & Export Data", ""),
        ("  5.7", "Fitur Paste Link Langsung", ""),
        ("  5.8", "Fitur Simpan & Muat Session", ""),
        ("BAB 6", "Workspace — Mengelola Data", ""),
        ("  6.1", "Membuka Workspace", ""),
        ("  6.2", "Memuat Data", ""),
        ("  6.3", "Pencarian Data", ""),
        ("  6.4", "Filter Lanjutan", ""),
        ("  6.5", "Mode Tampilan (Tabel, Kartu, Ringkasan)", ""),
        ("  6.6", "Export Data dari Workspace", ""),
        ("BAB 7", "Insight — Analisis & Visualisasi Data", ""),
        ("  7.1", "Membuka Halaman Insight", ""),
        ("  7.2", "Input Statistik Manual SIPEDE", ""),
        ("  7.3", "Upload & Analisis Surat Masuk", ""),
        ("  7.4", "Upload & Analisis Surat Keluar", ""),
        ("  7.5", "Pengelompokan Jenis Surat", ""),
        ("  7.6", "Pengelompokan Asal Surat", ""),
        ("  7.7", "Filter Bulan & Tahun", ""),
        ("  7.8", "Grafik & Visualisasi", ""),
        ("  7.9", "Export Dashboard sebagai Gambar", ""),
        ("BAB 8", "Tips & Solusi Masalah Umum", ""),
        ("  8.1", "Tips Penggunaan", ""),
        ("  8.2", "Masalah Umum & Solusinya", ""),
        ("  8.3", "FAQ (Pertanyaan yang Sering Diajukan)", ""),
    ]
    for num, title_text, _ in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title_text}")
        run.font.size = Pt(10)
        if num.startswith("BAB"):
            run.bold = True
            run.font.color.rgb = RGBColor(4, 78, 59)

    doc.add_page_break()

    # ================================================================
    #                    BAB 1: MENGENAL APLIKASI DASTA
    # ================================================================
    doc.add_heading('BAB 1 — Mengenal Aplikasi DASTA', level=1)

    doc.add_heading('1.1 Apa itu DASTA?', level=2)
    doc.add_paragraph(
        'DASTA (Deskripsi Data) adalah aplikasi website yang membantu Anda mengambil data '
        'secara otomatis (scraping) dari tiga sistem informasi Kejaksaan, yaitu SIPEDE, SPDP, '
        'dan DASTI. Selain scraping, DASTA juga menyediakan fitur untuk menganalisis, '
        'mengelola, dan mengekspor data yang telah dikumpulkan.'
    )
    doc.add_paragraph(
        'Dengan DASTA, pekerjaan yang sebelumnya membutuhkan input manual satu per satu '
        'dari tiga website berbeda kini dapat dilakukan secara otomatis dalam hitungan menit.'
    )

    doc.add_heading('1.2 Fitur Utama', level=2)
    features = [
        'Scraping Otomatis — Mengambil data dari SIPEDE, SPDP, dan DASTI secara otomatis.',
        'Monitoring Real-time — Memantau proses scraping, status server, dan progress secara langsung.',
        'Workspace — Mengelola, memfilter, dan mencari data yang sudah di-scrape.',
        'Insight — Menganalisis data surat masuk/keluar dengan grafik dan pengelompokan.',
        'Export Data — Mengunduh data dalam format Excel (.xlsx), CSV, atau JSON.',
        'Notifikasi Browser — Mendapat pemberitahuan otomatis ketika scraping selesai.',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('1.3 Struktur Menu Navigasi', level=2)
    doc.add_paragraph(
        'Setelah membuka aplikasi, Anda akan melihat menu navigasi di bagian atas halaman '
        'dengan struktur sebagai berikut:'
    )
    add_styled_table(doc,
        ['Menu', 'Sub-Menu', 'Penjelasan'],
        [
            ['Home', '—', 'Halaman utama berisi monitoring dan status semua scraper'],
            ['Scrapping ▾', 'SIPEDE', 'Halaman untuk scraping data dari sistem SIPEDE'],
            ['Scrapping ▾', 'SPDP', 'Halaman untuk scraping data dari sistem SPDP'],
            ['Scrapping ▾', 'DASTI', 'Halaman untuk scraping data dari sistem DASTI'],
            ['Workspace ▾', 'SIPEDE', 'Halaman untuk mengelola & menganalisis data SIPEDE'],
            ['Workspace ▾', 'SPDP', 'Halaman untuk mengelola & menganalisis data SPDP'],
            ['Insight', '—', 'Halaman analisis, visualisasi, dan statistik data surat'],
        ]
    )
    add_tip(doc, 'Menu "Scrapping" dan "Workspace" adalah dropdown. Arahkan kursor atau klik untuk melihat sub-menu.')

    doc.add_page_break()

    # ================================================================
    #                    BAB 2: HALAMAN UTAMA (HOME)
    # ================================================================
    doc.add_heading('BAB 2 — Halaman Utama (Home)', level=1)
    doc.add_paragraph(
        'Halaman Home adalah halaman pertama yang Anda lihat saat membuka aplikasi. '
        'Halaman ini berfungsi sebagai pusat monitoring dan kontrol untuk seluruh proses scraping.'
    )

    doc.add_heading('2.1 Tampilan Selamat Datang', level=2)
    doc.add_paragraph(
        'Di bagian atas halaman, Anda akan melihat tulisan "Selamat Datang di Dasta" '
        'beserta subtitle "Pusat monitoring & kontrol web scraping". '
        'Di sebelah kanan terdapat jam digital yang menunjukkan waktu dan tanggal saat ini '
        'dalam format Indonesia (contoh: "Senin, 11 Mar 2026").'
    )

    doc.add_heading('2.2 Kartu Status Ringkasan', level=2)
    doc.add_paragraph(
        'Di bawah header, terdapat empat kartu kecil yang menampilkan ringkasan status sistem:'
    )
    add_styled_table(doc,
        ['Kartu', 'Ikon', 'Penjelasan'],
        [
            ['Total Data', '📊', 'Jumlah total seluruh data yang berhasil di-scrape dari semua sumber'],
            ['Aktif', '⚡', 'Jumlah scraper yang sedang aktif berjalan saat ini'],
            ['Server', '🖥️', 'Jumlah server backend yang sedang online'],
            ['Update', '🕐', 'Waktu terakhir data diperbarui'],
        ]
    )

    doc.add_heading('2.3 Kartu Status Per Scraper', level=2)
    doc.add_paragraph(
        'Di bawah kartu ringkasan, terdapat tiga kartu besar untuk masing-masing scraper '
        '(SIPEDE, SPDP, dan DASTI). Setiap kartu menampilkan informasi berikut:'
    )
    info_items = [
        'Nama Scraper — Judul kartu (contoh: "SIPEDE Scraper").',
        'Status Badge — Menunjukkan kondisi scraper saat ini dengan warna:',
    ]
    for item in info_items:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_table(doc,
        ['Status', 'Warna', 'Artinya'],
        [
            ['Offline', 'Merah', 'Server backend tidak berjalan'],
            ['Standby', 'Abu-abu', 'Server aktif, browser belum dibuka'],
            ['Login', 'Oranye', 'Browser terbuka, menunggu login'],
            ['Ready', 'Biru', 'Siap untuk memulai scraping'],
            ['Scraping...', 'Hijau (berkedip)', 'Proses scraping sedang berjalan'],
            ['Done', 'Hijau', 'Scraping telah selesai'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Informasi lain di setiap kartu:')
    extra_info = [
        'Data — Jumlah data yang telah di-scrape.',
        'Halaman — Jumlah halaman yang telah diproses.',
        'Browser — Status browser (Open/Closed).',
        'Login — Status login (Yes/No).',
        'Terakhir di-scrape — Waktu terakhir scraping selesai.',
        'Tombol Export — Tombol untuk mengunduh data (Excel, CSV, JSON).',
    ]
    for item in extra_info:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.4 Status Server & Response Time', level=2)
    doc.add_paragraph(
        'Di bagian bawah halaman Home, terdapat bagian Status Server yang menampilkan '
        'kondisi koneksi ke setiap backend service:'
    )
    add_styled_table(doc,
        ['Server', 'Indikator', 'Keterangan'],
        [
            ['SIPEDE Backend (Port 5000)', 'Bulatan Hijau / Merah', 'Menunjukkan apakah server menyala atau mati'],
            ['SPDP Backend (Port 5001)', 'Bulatan Hijau / Merah', 'Menunjukkan koneksi ke server SPDP'],
            ['DASTI Backend (Port 5002)', 'Bulatan Hijau / Merah', 'Menunjukkan koneksi ke server DASTI'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Response time juga ditampilkan untuk membantu mengidentifikasi masalah jaringan. '
        'Jika response time tinggi atau server offline, periksa apakah backend sudah dijalankan.'
    )

    doc.add_heading('2.5 Activity Log', level=2)
    doc.add_paragraph(
        'Activity Log menampilkan riwayat aktivitas scraping secara kronologis. '
        'Setiap log memiliki level warna yang berbeda:'
    )
    add_styled_table(doc,
        ['Level', 'Warna', 'Contoh Pesan'],
        [
            ['Info', 'Biru', '"Scraping dimulai", "Browser dibuka"'],
            ['Success', 'Hijau', '"Scraping selesai — 150 data berhasil diambil"'],
            ['Warning', 'Kuning', '"Pagination tidak terdeteksi"'],
            ['Error', 'Merah', '"Connection failed — server offline"'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Log dapat dihapus dengan mengklik tombol "Hapus Log" di bagian atas daftar log.')

    doc.add_heading('2.6 Fitur Auto-Refresh (LIVE)', level=2)
    doc.add_paragraph(
        'Di pojok kanan atas halaman Home terdapat tombol LIVE/PAUSED:'
    )
    live_info = [
        'LIVE (hijau) — Data di halaman akan diperbarui otomatis setiap 5 detik.',
        'PAUSED (abu-abu) — Pembaruan otomatis dimatikan.',
        'Klik tombol untuk mengaktifkan atau menonaktifkan auto-refresh.',
    ]
    for item in live_info:
        doc.add_paragraph(item, style='List Bullet')
    add_tip(doc, 'Aktifkan mode LIVE saat ingin memantau proses scraping secara real-time.')

    doc.add_page_break()

    # ================================================================
    #                    BAB 3: SCRAPING SIPEDE
    # ================================================================
    doc.add_heading('BAB 3 — Scraping SIPEDE', level=1)
    doc.add_paragraph(
        'Bab ini menjelaskan langkah-langkah lengkap untuk melakukan scraping data dari '
        'sistem SIPEDE (Sistem Informasi Persuratan Desa). SIPEDE Scraper mengambil data '
        'surat terkirim dari website SIPEDE Kejaksaan.'
    )

    doc.add_heading('3.1 Persiapan Sebelum Scraping', level=2)
    doc.add_paragraph('Pastikan hal-hal berikut sebelum memulai:')
    prep = [
        'Koneksi internet aktif dan stabil.',
        'Anda memiliki akun (username & password) untuk login ke SIPEDE.',
        'Server backend SIPEDE sudah berjalan (cek di halaman Home > Status Server).',
        'Google Chrome terinstal di komputer.',
    ]
    for item in prep:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Langkah 1 — Buka Browser', level=2)
    add_numbered_steps(doc, [
        'Klik menu "Scrapping" di navigasi atas.',
        'Pilih "SIPEDE" dari dropdown menu.',
        'Anda akan melihat halaman SIPEDE Scraper dengan judul "Mulai Scraping SIPEDE".',
        'Klik tombol hijau "Buka Browser SIPEDE".',
        'Tunggu hingga browser Chrome terbuka secara otomatis dan halaman login SIPEDE tampil.',
        'Tombol akan berubah menjadi "Membuka browser..." selama proses loading.',
    ])
    add_important(doc,
        'Jangan menutup browser Chrome yang dibuka oleh aplikasi. '
        'Browser ini dikontrol oleh sistem DASTA untuk proses scraping.')

    doc.add_heading('3.3 Langkah 2 — Login ke SIPEDE', level=2)
    doc.add_paragraph(
        'Setelah browser terbuka, Anda perlu login secara manual:'
    )
    add_numbered_steps(doc, [
        'Di browser Chrome yang terbuka, masukkan username Anda.',
        'Masukkan password Anda.',
        'Selesaikan CAPTCHA jika diminta.',
        'Klik tombol login di website SIPEDE.',
        'Tunggu hingga login berhasil dan halaman utama SIPEDE tampil.',
        'Kembali ke aplikasi DASTA, klik tombol "Saya Sudah Login".',
        'Sistem akan memverifikasi login Anda. Jika berhasil, Anda akan lanjut ke langkah berikutnya.',
    ])
    add_tip(doc,
        'Setelah login berhasil, sistem otomatis mengarahkan browser ke halaman "Surat Terkirim".')

    doc.add_heading('3.4 Langkah 3 — Konfigurasi & Mulai Scraping', level=2)
    doc.add_paragraph(
        'Setelah login terverifikasi, Anda akan melihat panel konfigurasi "Siap Scraping!" '
        'dengan beberapa pengaturan:'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('A. Pilih Tahun Data')
    run.bold = True
    doc.add_paragraph(
        'Gunakan dropdown "Pilih Tahun Data:" untuk memilih tahun data yang ingin di-scrape. '
        'Tahun yang tersedia mulai dari 2022 hingga tahun berjalan.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('B. Atur Entries Per Page')
    run.bold = True
    doc.add_paragraph(
        'Pilih jumlah data per halaman dari dropdown "Entries Per Page:". '
        'Pilihan yang tersedia: 10, 25, 50, atau 100. '
        'Semakin banyak entries per page, semakin sedikit halaman yang perlu di-scrape.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('C. Maksimum Halaman')
    run.bold = True
    doc.add_paragraph(
        'Masukkan angka pada kolom "Maksimum halaman (0 = semua):". '
        'Isi 0 untuk scraping semua halaman, atau masukkan angka tertentu '
        'untuk membatasi jumlah halaman.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('D. Informasi Tabel')
    run.bold = True
    doc.add_paragraph(
        'Sebelum memulai, Anda dapat melihat preview informasi tabel:'
    )
    info_items = [
        'Total Data — Jumlah total data yang tersedia di website.',
        'Total Halaman — Jumlah halaman data.',
        'Kolom — Daftar nama kolom yang akan diambil.',
    ]
    for item in info_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('E. Mulai Scraping')
    run.bold = True
    doc.add_paragraph(
        'Klik tombol hijau "▶ Mulai Scraping" untuk memulai proses. '
        'Anda juga dapat klik "Refresh Data" untuk memperbarui informasi tabel.'
    )

    doc.add_heading('3.5 Langkah 4 — Melihat Progress Scraping', level=2)
    doc.add_paragraph(
        'Selama proses scraping, aplikasi menampilkan informasi progress secara real-time:'
    )
    add_styled_table(doc,
        ['Informasi', 'Keterangan'],
        [
            ['Progress Bar', 'Bar visual yang menunjukkan persentase penyelesaian'],
            ['Page X / Y', 'Halaman yang sedang di-scrape vs total halaman'],
            ['Halaman Saat Ini', 'Nomor halaman yang sedang diproses'],
            ['Halaman Selesai', 'Jumlah halaman yang sudah berhasil di-scrape'],
            ['Data', 'Jumlah data yang sudah berhasil diambil'],
            ['Waktu', 'Durasi waktu scraping (dalam detik)'],
            ['Status Detail', '"Navigasi...", "Menunggu...", atau "Sedang scraping..."'],
        ]
    )
    add_warning(doc, 'Jangan menutup browser atau mematikan aplikasi selama proses scraping berlangsung.')

    doc.add_heading('3.6 Langkah 5 — Hasil Scraping & Export Data', level=2)
    doc.add_paragraph(
        'Ketika proses selesai, Anda akan melihat halaman "✅ Scraping Selesai!" '
        'yang menampilkan ringkasan dan data hasil scraping:'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ringkasan Hasil:')
    run.bold = True
    summary = [
        'Halaman — Jumlah halaman yang berhasil di-scrape.',
        'Total Data — Jumlah keseluruhan baris data yang diperoleh.',
        'Waktu — Total durasi proses scraping.',
    ]
    for item in summary:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Melihat Data:')
    run.bold = True
    doc.add_paragraph(
        'Data ditampilkan dalam tabel dengan fitur pencarian ("Cari data...") dan paginasi. '
        'Gunakan tombol "← Prev" dan "Next →" untuk berpindah halaman data.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Export Data:')
    run.bold = True
    doc.add_paragraph('Klik salah satu tombol export untuk mengunduh data:')
    export_options = [
        'Tombol "Excel" — Mengunduh file .xlsx yang dapat dibuka di Microsoft Excel.',
        'Tombol "CSV" — Mengunduh file .csv untuk kompatibilitas dengan berbagai aplikasi.',
        'Tombol "JSON" — Mengunduh file .json untuk keperluan teknis.',
    ]
    for opt in export_options:
        doc.add_paragraph(opt, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Selesai:')
    run.bold = True
    doc.add_paragraph(
        'Klik tombol merah "Selesai & Tutup Browser" untuk menutup browser dan mengakhiri sesi scraping.'
    )

    doc.add_heading('3.7 Scraping Tahun Lain', level=2)
    doc.add_paragraph(
        'Jika ingin scraping data tahun lain tanpa menutup browser:'
    )
    add_numbered_steps(doc, [
        'Di halaman hasil, cari dropdown pilihan tahun di bagian bawah.',
        'Pilih tahun yang berbeda dari dropdown.',
        'Klik tombol "Scraping Tahun Lain".',
        'Proses scraping akan dimulai kembali untuk tahun yang dipilih.',
    ])

    doc.add_page_break()

    # ================================================================
    #                    BAB 4: SCRAPING SPDP
    # ================================================================
    doc.add_heading('BAB 4 — Scraping SPDP', level=1)
    doc.add_paragraph(
        'Bab ini menjelaskan cara melakukan scraping data dari sistem SPDP '
        '(Surat Pemberitahuan Dimulainya Penyidikan). SPDP Scraper diakses melalui '
        'jaringan LAN (lokal) kantor.'
    )

    doc.add_heading('4.1 Persiapan Sebelum Scraping', level=2)
    doc.add_paragraph('Pastikan hal-hal berikut sebelum memulai:')
    prep = [
        'Komputer terhubung ke jaringan LAN kantor (kabel LAN atau WiFi kantor).',
        'Server SPDP aktif dan dapat diakses di alamat 10.35.0.101:4111.',
        'Anda memiliki akun (username & password) untuk login ke SPDP.',
        'Server backend SPDP sudah berjalan (cek di halaman Home > Status Server).',
    ]
    for item in prep:
        doc.add_paragraph(item, style='List Bullet')
    add_warning(doc,
        'SPDP hanya dapat diakses melalui jaringan LAN. '
        'Scraping tidak akan berfungsi jika menggunakan internet publik.')

    doc.add_heading('4.2 Langkah 1 — Buka Browser', level=2)
    add_numbered_steps(doc, [
        'Klik menu "Scrapping" di navigasi atas.',
        'Pilih "SPDP" dari dropdown menu.',
        'Anda akan melihat halaman SPDP Scraper.',
        'Klik tombol "Buka Browser SPDP".',
        'Browser Chrome akan terbuka ke halaman login SPDP.',
    ])
    add_tip(doc,
        'Jika muncul pesan error koneksi, pastikan kabel LAN terhubung atau WiFi kantor aktif.')

    doc.add_heading('4.3 Langkah 2 — Login & Navigasi ke Halaman Data', level=2)
    add_numbered_steps(doc, [
        'Di browser Chrome yang terbuka, login menggunakan username dan password SPDP Anda.',
        'Setelah login, navigasi ke halaman yang berisi data tabel yang ingin di-scrape.',
        'Pastikan TABEL DATA sudah tampil di layar browser.',
        'Di bagian bawah halaman scraper, Anda dapat melihat "URL saat ini:" yang menunjukkan URL halaman browser.',
        'Kembali ke aplikasi DASTA, klik tombol "Saya Sudah di Halaman Data".',
        'Sistem akan mencoba mendeteksi tabel secara otomatis.',
    ])
    add_tip(doc,
        'Jika Anda mengubah filter atau navigasi di website SPDP, '
        'klik "Refresh Data" untuk memperbarui informasi tabel di DASTA.')

    doc.add_heading('4.4 Langkah 3 — Deteksi Tabel & Konfigurasi', level=2)
    doc.add_paragraph(
        'Setelah tabel terdeteksi, halaman akan menampilkan informasi "Tabel Terdeteksi!" '
        'dengan detail berikut:'
    )
    info_items = [
        'Total Data — Jumlah baris data yang terdeteksi.',
        'Total Halaman — Jumlah halaman pagination.',
        'Kolom — Daftar nama kolom yang akan diambil (ditampilkan sebagai tag).',
    ]
    for item in info_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Konfigurasi Scraping:')
    run.bold = True
    doc.add_paragraph()
    config_items = [
        'Rentang Halaman — Isi "Dari:" dan "Sampai:" untuk menentukan halaman mana yang akan di-scrape. Isi 0 pada "Sampai:" untuk scraping sampai halaman terakhir.',
        'Filter Tahun — Pilih tahun dari dropdown. Pilih "Semua Tahun" untuk mengambil data semua tahun.',
    ]
    for item in config_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Tombol yang tersedia:')
    buttons = [
        '"Pilih Data Lain" — Kembali ke langkah navigasi untuk memilih tabel data yang berbeda.',
        '"Refresh Data" — Memperbarui informasi tabel jika ada perubahan di website.',
        '"▶ Mulai Scraping" — Memulai proses scraping sesuai konfigurasi.',
    ]
    for b in buttons:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('4.5 Langkah 4 — Melihat Progress Scraping', level=2)
    doc.add_paragraph(
        'Seperti pada SIPEDE, progress scraping ditampilkan secara real-time dengan '
        'progress bar, informasi halaman saat ini, data yang sudah diambil, dan waktu elapsed. '
        'Notifikasi browser akan muncul saat proses selesai.'
    )

    doc.add_heading('4.6 Langkah 5 — Hasil Scraping & Export Data', level=2)
    doc.add_paragraph(
        'Setelah selesai, halaman "✅ Scraping SPDP Selesai!" akan tampil dengan:'
    )
    results = [
        'Ringkasan hasil (halaman, total data, waktu).',
        'Tabel data dengan pencarian dan paginasi.',
        'Tombol export: Excel, CSV, JSON.',
        'Tombol "Scraping Data Baru" untuk memulai scraping baru.',
        'Tombol "Hapus Data" untuk menghapus data scraping yang telah tersimpan.',
        'Tombol "Selesai & Tutup Browser" untuk mengakhiri sesi.',
    ]
    for item in results:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    #                    BAB 5: SCRAPING DASTI
    # ================================================================
    doc.add_heading('BAB 5 — Scraping DASTI', level=1)
    doc.add_paragraph(
        'Bab ini menjelaskan cara melakukan scraping data dari sistem DASTI. '
        'DASTI Scraper memiliki alur yang sedikit lebih kompleks karena website DASTI '
        'dilengkapi dengan captcha dan memerlukan navigasi multi-level.'
    )

    doc.add_heading('5.1 Persiapan Sebelum Scraping', level=2)
    doc.add_paragraph('Pastikan hal-hal berikut sebelum memulai:')
    prep = [
        'Koneksi internet aktif dan stabil.',
        'Anda memiliki akun (username & password) untuk login ke website DASTI.',
        'Server backend DASTI sudah berjalan (cek di halaman Home > Status Server).',
        'Google Chrome terinstal dan versi terbaru.',
    ]
    for item in prep:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2 Langkah 1 — Buka Browser', level=2)
    add_numbered_steps(doc, [
        'Klik menu "Scrapping" di navigasi atas.',
        'Pilih "DASTI" dari dropdown menu.',
        'Anda akan melihat halaman DASTI Scraper.',
        'Klik tombol "Open DASTI Login".',
        'Browser Chrome akan terbuka ke halaman login DASTI.',
    ])

    doc.add_heading('5.3 Langkah 2 — Login ke DASTI (dengan Captcha)', level=2)
    doc.add_paragraph(
        'Login DASTI memerlukan penyelesaian captcha secara manual:'
    )
    add_numbered_steps(doc, [
        'Di browser Chrome, masukkan username dan password DASTI Anda.',
        'Jika captcha muncul, aplikasi DASTA akan menampilkan peringatan "🛡️ Captcha terdeteksi!".',
        'Selesaikan captcha secara manual di browser (bisa berupa reCAPTCHA, gambar, atau teks).',
        'Klik tombol login di browser DASTI.',
        'Tunggu hingga login berhasil.',
        'Kembali ke DASTA, klik tombol "Saya Sudah Login".',
        'Jika verifikasi berhasil, pesan "Login Berhasil!" akan muncul.',
    ])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Alternatif — Load Session Lama:')
    run.bold = True
    doc.add_paragraph(
        'Jika Anda pernah login sebelumnya dan menyimpan session, klik tombol "Load Session Lama" '
        'untuk memuat session yang tersimpan tanpa perlu login ulang. '
        'Session berlaku selama 24 jam sejak terakhir disimpan.'
    )

    doc.add_heading('5.4 Langkah 3 — Navigasi ke Halaman Data', level=2)
    doc.add_paragraph(
        'Setelah login berhasil, Anda perlu menuju halaman yang berisi tabel data. '
        'Ada dua cara untuk melakukan ini:'
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('📋 Opsi 1: Paste Link Langsung (Direkomendasikan)')
    run.bold = True
    run.font.color.rgb = RGBColor(4, 78, 59)
    add_numbered_steps(doc, [
        'Salin (copy) URL halaman data DASTI yang ingin di-scrape.',
        'Tempel (paste) URL tersebut ke kolom "URL Halaman Data Tabel:" di aplikasi.',
        'Klik tombol "🔗 Navigasi ke URL".',
        'Browser akan otomatis berpindah ke halaman tersebut.',
        'Sistem akan mencoba mendeteksi tabel secara otomatis.',
    ])
    add_tip(doc, 'Metode Paste Link jauh lebih cepat (5-10 detik) dibandingkan navigasi manual (30-60 detik).')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('🖱️ Opsi 2: Navigasi Manual')
    run.bold = True
    run.font.color.rgb = RGBColor(4, 78, 59)
    add_numbered_steps(doc, [
        'Navigasi secara manual di browser DASTI hingga menemukan halaman yang berisi tabel data.',
        'Kembali ke DASTA, klik tombol "🔍 Deteksi Tabel".',
        'Sistem akan mendeteksi struktur tabel pada halaman.',
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('💾 Simpan Session (Opsional):')
    run.bold = True
    doc.add_paragraph(
        'Klik tombol "💾 Simpan Session" untuk menyimpan session login Anda. '
        'Session ini bisa digunakan kembali di sesi berikutnya.'
    )

    doc.add_heading('5.5 Langkah 4 — Konfigurasi & Mulai Scraping', level=2)
    doc.add_paragraph(
        'Setelah tabel terdeteksi, halaman "Siap Scraping!" akan tampil. Konfigurasi yang tersedia:'
    )
    config = [
        'Halaman awal — Nomor halaman pertama yang akan di-scrape (default: 1).',
        'Halaman akhir — Nomor halaman terakhir (isi 0 untuk scraping semua halaman).',
    ]
    for item in config:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Informasi yang ditampilkan:')
    info = [
        'Total Data — Jumlah total data di tabel.',
        'Total Halaman — Jumlah halaman pagination.',
        'Kolom — Daftar kolom yang akan diambil.',
    ]
    for item in info:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph(
        'Klik "Refresh Tabel" jika perlu memperbarui, atau klik "▶ Mulai Scraping" untuk memulai.'
    )

    doc.add_heading('5.6 Langkah 5 — Hasil Scraping & Export Data', level=2)
    doc.add_paragraph(
        'Setelah proses selesai, halaman "✅ Scraping Selesai!" akan tampil dengan:'
    )
    results = [
        'Ringkasan: halaman yang di-scrape, total data, dan waktu.',
        'Tabel data hasil dengan pencarian dan paginasi.',
        'Tombol export: CSV, Excel, dan JSON.',
        'Tombol "Scraping Baru" untuk memulai sesi scraping baru.',
    ]
    for item in results:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.7 Fitur Paste Link Langsung', level=2)
    doc.add_paragraph(
        'Fitur Paste Link adalah cara tercepat untuk mengarahkan scraper ke halaman data '
        'target. Berikut keuntungannya:'
    )
    add_styled_table(doc,
        ['Keuntungan', 'Penjelasan'],
        [
            ['Lebih Cepat', 'Hanya 5-10 detik vs 30-60 detik navigasi manual'],
            ['Lebih Akurat', 'URL langsung menuju halaman data spesifik'],
            ['Dapat Diulang', 'URL bisa disimpan dan dipakai lagi untuk scraping berikutnya'],
            ['Melewati Navigasi', 'Tidak perlu klik menu demi menu di website DASTI'],
        ]
    )

    doc.add_heading('5.8 Fitur Simpan & Muat Session', level=2)
    doc.add_paragraph(
        'DASTI Scraper mendukung fitur penyimpanan session untuk menghindari login berulang:'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cara Menyimpan Session:')
    run.bold = True
    add_numbered_steps(doc, [
        'Login ke DASTI seperti biasa.',
        'Setelah login berhasil, klik tombol "💾 Simpan Session".',
        'Session (cookies) akan disimpan ke database lokal.',
        'Pesan konfirmasi akan muncul jika berhasil.',
    ])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cara Memuat Session:')
    run.bold = True
    add_numbered_steps(doc, [
        'Di halaman login DASTI, klik tombol "Load Session Lama".',
        'Sistem akan memuat session yang tersimpan.',
        'Jika session masih valid, Anda akan langsung lanjut tanpa login.',
        'Jika session sudah expired (>24 jam), Anda perlu login ulang.',
    ])
    add_important(doc,
        'Aplikasi TIDAK menyimpan password Anda. '
        'Hanya cookies session yang disimpan untuk kemudahan penggunaan.')

    doc.add_page_break()

    # ================================================================
    #                    BAB 6: WORKSPACE
    # ================================================================
    doc.add_heading('BAB 6 — Workspace: Mengelola Data', level=1)
    doc.add_paragraph(
        'Workspace adalah fitur untuk mengelola, memfilter, dan menganalisis data '
        'yang telah di-scrape. Terdapat dua workspace: SIPEDE dan SPDP.'
    )

    doc.add_heading('6.1 Membuka Workspace', level=2)
    add_numbered_steps(doc, [
        'Klik menu "Workspace" di navigasi atas.',
        'Pilih "SIPEDE" atau "SPDP" dari dropdown menu.',
        'Halaman workspace akan terbuka.',
    ])

    doc.add_heading('6.2 Memuat Data', level=2)
    doc.add_paragraph('Ada dua cara untuk memuat data ke workspace:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cara 1: Muat dari API (Hasil Scraping)')
    run.bold = True
    add_numbered_steps(doc, [
        'Klik tombol "Muat dari API SIPEDE" (atau "Muat dari API SPDP").',
        'Sistem akan mengambil data dari database backend.',
        'Data akan ditampilkan dalam tabel.',
        'Pesan "Berhasil memuat X data" akan muncul.',
    ])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cara 2: Upload File Excel')
    run.bold = True
    add_numbered_steps(doc, [
        'Seret (drag) file Excel (.xlsx atau .xls) ke area upload, ATAU',
        'Klik area upload untuk membuka dialog pemilihan file.',
        'Pilih file Excel yang ingin dimuat.',
        'Data dari Excel akan diparse dan ditampilkan di tabel.',
    ])

    doc.add_heading('6.3 Pencarian Data', level=2)
    add_numbered_steps(doc, [
        'Temukan kolom pencarian "Cari data..." di bagian atas tabel.',
        'Ketikkan kata kunci yang ingin dicari.',
        'Klik ikon 🔍 atau tekan Enter untuk memulai pencarian.',
        'Tabel akan menampilkan hanya data yang sesuai dengan kata kunci.',
        'Hapus teks pencarian untuk menampilkan semua data kembali.',
    ])

    doc.add_heading('6.4 Filter Lanjutan', level=2)
    doc.add_paragraph(
        'Untuk filtering yang lebih spesifik, gunakan fitur Filter Lanjutan:'
    )
    add_numbered_steps(doc, [
        'Klik ikon Filter (🔽) di toolbar.',
        'Panel filter akan terbuka.',
        'Pilih kolom yang ingin difilter dari dropdown "Pilih kolom".',
        'Pilih operator filter:',
    ])
    add_styled_table(doc,
        ['Operator', 'Arti', 'Contoh'],
        [
            ['contains', 'Mengandung teks', '"surat" → menampilkan semua data yang mengandung kata "surat"'],
            ['equals', 'Sama persis', '"2025" → hanya data yang bernilai tepat "2025"'],
            ['startsWith', 'Diawali dengan', '"SP" → data yang diawali "SP..."'],
            ['endsWith', 'Diakhiri dengan', '"RI" → data yang diakhiri "...RI"'],
            ['notContains', 'Tidak mengandung', '"draft" → data yang tidak mengandung "draft"'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        '5. Masukkan nilai filter pada kolom input.\n'
        '6. Filter akan langsung diterapkan pada tabel.'
    )

    doc.add_heading('6.5 Mode Tampilan (Tabel, Kartu, Ringkasan)', level=2)
    doc.add_paragraph('Workspace menyediakan tiga mode tampilan data:')
    add_styled_table(doc,
        ['Mode', 'Ikon', 'Penjelasan'],
        [
            ['Tabel', '📊', 'Mode default — Data ditampilkan dalam format tabel dengan baris dan kolom'],
            ['Kartu (Cards)', '🃏', 'Data ditampilkan sebagai kartu individual untuk setiap baris data'],
            ['Ringkasan (Summary)', '📋', 'Menampilkan statistik ringkasan: jumlah data unik, nilai teratas per kolom'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Klik nama mode di toolbar untuk beralih tampilan.')

    doc.add_heading('6.6 Export Data dari Workspace', level=2)
    doc.add_paragraph('Data di workspace dapat diekspor dalam beberapa format:')
    add_styled_table(doc,
        ['Tombol', 'Format', 'Keterangan'],
        [
            ['"Export ke Excel"', '.xlsx', 'File spreadsheet Excel — format paling umum'],
            ['"Export ke CSV"', '.csv', 'Comma-Separated Values — kompatibel universal'],
            ['"Copy ke Clipboard"', 'Teks', 'Menyalin data ke clipboard untuk paste di aplikasi lain'],
        ]
    )
    add_tip(doc, 'Format Excel direkomendasikan karena mendukung format tabel dan kolom yang lebih rapi.')

    doc.add_page_break()

    # ================================================================
    #                    BAB 7: INSIGHT
    # ================================================================
    doc.add_heading('BAB 7 — Insight: Analisis & Visualisasi Data', level=1)
    doc.add_paragraph(
        'Halaman Insight adalah fitur analisis data lanjutan yang memungkinkan Anda '
        'menganalisis data surat masuk dan surat keluar dengan pengelompokan, '
        'visualisasi grafik, dan statistik.'
    )

    doc.add_heading('7.1 Membuka Halaman Insight', level=2)
    add_numbered_steps(doc, [
        'Klik menu "Insight" di navigasi atas (ikon kubus).',
        'Halaman Insight akan terbuka dengan beberapa bagian/section.',
    ])

    doc.add_heading('7.2 Input Statistik Manual SIPEDE', level=2)
    doc.add_paragraph(
        'Bagian ini memungkinkan Anda memasukkan data statistik SIPEDE '
        'yang tidak tersedia melalui scraping otomatis:'
    )
    add_styled_table(doc,
        ['Kategori', 'Input Yang Diperlukan'],
        [
            ['Status Desa', 'Jumlah Status Aktif dan Status Tidak Aktif'],
            ['Registrasi SIPEDE', 'Jumlah Tercatat di SIPEDE dan Tidak Tercatat di SIPEDE'],
            ['E-Signature', 'Jumlah Terdaftar E-sign dan Tidak Terdaftar E-sign'],
        ]
    )
    doc.add_paragraph()
    add_numbered_steps(doc, [
        'Masukkan angka di setiap kolom input.',
        'Klik tombol "Simpan Statistik Manual".',
        'Pesan "Statistik berhasil disimpan!" akan muncul.',
        'Data akan tersimpan di browser (localStorage) dan tampil sebagai Donut Chart.',
    ])
    add_tip(doc, 'Data statistik manual akan tetap tersimpan meskipun browser ditutup karena menggunakan localStorage.')

    doc.add_heading('7.3 Upload & Analisis Surat Masuk', level=2)
    doc.add_paragraph(
        'Fitur ini menganalisis data surat masuk dari file Excel yang Anda upload:'
    )
    add_numbered_steps(doc, [
        'Klik tombol "Pilih File Excel Surat Masuk".',
        'Pilih file Excel (.xlsx atau .xls) yang berisi data surat masuk.',
        'Sistem akan membaca dan menganalisis file secara otomatis.',
        'Informasi ringkasan file akan ditampilkan:',
    ])
    info_items = [
        'Total Baris — Jumlah baris data dalam file.',
        'Jenis Surat — Jumlah jenis surat yang ditemukan.',
        'Jenis Cocok — Jumlah jenis surat yang berhasil dikategorikan.',
        'Jenis Belum Cocok — Jumlah jenis surat yang belum dikategorikan.',
        'Asal — Jumlah asal surat yang ditemukan.',
        'Asal Cocok / Belum Cocok — Status pengelompokan asal surat.',
    ]
    for item in info_items:
        doc.add_paragraph(f'   • {item}')

    doc.add_heading('7.4 Upload & Analisis Surat Keluar', level=2)
    doc.add_paragraph(
        'Fitur ini sama dengan Surat Masuk, tetapi untuk data surat keluar. '
        'Upload file Excel surat keluar dan sistem akan menganalisisnya '
        'dengan cara yang sama.'
    )

    doc.add_heading('7.5 Pengelompokan Jenis Surat', level=2)
    doc.add_paragraph(
        'Setelah upload data surat, klik tab "Pengelompokan Jenis Surat" '
        'untuk mengelompokkan data berdasarkan jenis/tipe surat.'
    )
    doc.add_paragraph('Terdapat 7 kategori standar:')
    add_styled_table(doc,
        ['No', 'Kategori', 'Penjelasan'],
        [
            ['1', 'Surat Masuk Biasa', 'Surat masuk dengan klasifikasi umum/biasa'],
            ['2', 'Surat Masuk Rahasia', 'Surat masuk dengan klasifikasi rahasia/terbatas'],
            ['3', 'Keputusan / SP / ST', 'Surat keputusan, surat perintah, surat tugas'],
            ['4', 'Nota Dinas / Memo', 'Nota dinas internal dan memorandum'],
            ['5', 'Laporan', 'Laporan-laporan resmi'],
            ['6', 'Email / Memorandum', 'Komunikasi elektronik dan memorandum'],
            ['7', 'Surat Keluar', 'Surat yang dikirim keluar organisasi'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Cara mengelompokkan jenis surat:')
    add_numbered_steps(doc, [
        'Setiap item surat ditampilkan dengan status ✓ (Sudah) atau ✗ (Belum) dikategorikan.',
        'Klik item yang belum dikategorikan.',
        'Pilih kategori yang sesuai dari dialog yang muncul.',
        'Klik "Simpan" untuk menyimpan pengelompokan.',
        'Gunakan kolom pencarian "Cari dalam kategori..." untuk memfilter item.',
    ])

    doc.add_heading('7.6 Pengelompokan Asal Surat', level=2)
    doc.add_paragraph(
        'Klik tab "Pengelompokan Asal" untuk mengelompokkan surat berdasarkan '
        'asal/pengirim surat. Fitur ini menyediakan:'
    )
    features = [
        'Kelompok bawaan (8+ grup standar) untuk pengelompokan umum.',
        'Fitur membuat kelompok kustom — Anda dapat menambahkan grup baru sesuai kebutuhan.',
        'Assignment per item — Setiap surat dapat di-assign ke kelompok tertentu.',
        'Pencarian dan filter — Mencari item dalam pengelompokan.',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('7.7 Filter Bulan & Tahun', level=2)
    doc.add_paragraph('Fitur filter untuk mempersempit data yang dianalisis:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Filter Bulan:')
    run.bold = True
    doc.add_paragraph(
        'Gunakan dropdown "Dari Bulan:" dan "Sampai Bulan:" untuk memfilter data '
        'berdasarkan rentang bulan (1-12).'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Override Tahun:')
    run.bold = True
    doc.add_paragraph(
        'Gunakan kolom "Override Tahun:" untuk mengganti tahun data pada file Excel. '
        'Berguna jika format tahun di file berbeda dari yang diinginkan.'
    )

    doc.add_heading('7.8 Grafik & Visualisasi', level=2)
    doc.add_paragraph('Halaman Insight menampilkan beberapa jenis grafik:')
    add_styled_table(doc,
        ['Jenis Grafik', 'Data yang Ditampilkan'],
        [
            ['Donut Chart (3 buah)', 'Distribusi Status Desa, Registrasi SIPEDE, dan E-sign'],
            ['Bar Chart — Tren Frekuensi', 'Tren jumlah surat masuk per bulan/periode'],
            ['Tabel Jenis Surat', 'Distribusi jenis surat per kategori dalam format tabel'],
            ['Tabel Asal Surat', 'Distribusi asal surat masuk dalam format tabel'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Setiap grafik memiliki judul, label, dan legenda warna untuk memudahkan pembacaan.'
    )

    doc.add_heading('7.9 Export Dashboard sebagai Gambar', level=2)
    doc.add_paragraph(
        'Anda dapat menyimpan tampilan dashboard/grafik Insight sebagai gambar:'
    )
    add_numbered_steps(doc, [
        'Pastikan grafik yang ingin disimpan sudah tampil di layar.',
        'Klik tombol screenshot/download yang tersedia di area grafik.',
        'Sistem akan men-capture tampilan menggunakan html2canvas.',
        'File gambar akan otomatis terunduh ke komputer Anda.',
    ])
    add_tip(doc, 'Gunakan fitur ini untuk menyisipkan grafik ke dalam laporan atau presentasi.')

    doc.add_page_break()

    # ================================================================
    #                    BAB 8: TIPS & SOLUSI MASALAH
    # ================================================================
    doc.add_heading('BAB 8 — Tips & Solusi Masalah Umum', level=1)

    doc.add_heading('8.1 Tips Penggunaan', level=2)

    tips = [
        ('Gunakan Paste Link untuk DASTI',
         'Salin URL halaman data DASTI dan tempel langsung di DASTA alih-alih '
         'navigasi manual. Ini menghemat waktu hingga 50 detik.'),
        ('Simpan Session DASTI',
         'Setelah login DASTI berhasil, selalu klik "Simpan Session" agar bisa '
         'menggunakan "Load Session Lama" di sesi berikutnya tanpa login ulang.'),
        ('Aktifkan Mode LIVE',
         'Saat memantau proses scraping di halaman Home, aktifkan mode LIVE '
         'untuk melihat update otomatis setiap 5 detik.'),
        ('Pilih Entries Per Page 100 untuk SIPEDE',
         'Menggunakan 100 entries per page mengurangi jumlah halaman yang perlu '
         'di-scrape, sehingga proses lebih cepat.'),
        ('Cek Status Server Terlebih Dahulu',
         'Sebelum memulai scraping, pastikan semua server menunjukkan status '
         '"Online" (hijau) di halaman Home.'),
        ('Gunakan Filter Workspace',
         'Gunakan kombinasi filter lanjutan (contains, startsWith, dll.) '
         'di Workspace untuk menemukan data spesifik dengan cepat.'),
        ('Export ke Excel',
         'Format Excel (.xlsx) direkomendasikan untuk export karena mendukung '
         'format tabel, kolom, dan pemfilteran di Microsoft Excel.'),
        ('Izinkan Notifikasi Browser',
         'Saat diminta izin notifikasi oleh browser, pilih "Allow" agar Anda '
         'mendapat pemberitahuan otomatis ketika scraping selesai.'),
    ]

    for title_text, desc in tips:
        p = doc.add_paragraph()
        run = p.add_run(f'💡 {title_text}')
        run.bold = True
        run.font.color.rgb = RGBColor(4, 78, 59)
        doc.add_paragraph(desc)
        doc.add_paragraph()

    doc.add_heading('8.2 Masalah Umum & Solusinya', level=2)

    problems = [
        {
            'masalah': 'Browser Chrome tidak terbuka saat klik "Buka Browser"',
            'penyebab': 'ChromeDriver tidak sesuai versi Chrome, atau Google Chrome belum terinstal.',
            'solusi': [
                'Perbarui Google Chrome ke versi terbaru.',
                'Jalankan file Fix-DastiChromeDriver.bat.',
                'Pastikan server backend sudah berjalan (cek Status Server di halaman Home).',
            ]
        },
        {
            'masalah': 'Server menunjukkan status "Offline" (merah) di halaman Home',
            'penyebab': 'Backend service belum dijalankan atau port sedang digunakan aplikasi lain.',
            'solusi': [
                'Jalankan ulang aplikasi menggunakan Start-WebScraper.bat.',
                'Atau jalankan backend secara manual di terminal terpisah.',
                'Cek apakah port (5000/5001/5002) digunakan oleh aplikasi lain.',
            ]
        },
        {
            'masalah': 'Tidak bisa mengakses SPDP (Connection Refused)',
            'penyebab': 'Komputer tidak terhubung ke jaringan LAN.',
            'solusi': [
                'Pastikan kabel LAN terhubung atau WiFi LAN kantor aktif.',
                'Coba ping 10.35.0.101 dari Command Prompt.',
                'Hubungi administrator jaringan jika masalah berlanjut.',
            ]
        },
        {
            'masalah': 'Captcha DASTI tidak bisa diselesaikan / login gagal',
            'penyebab': 'Halaman belum dimuat sempurna atau captcha berubah tipe.',
            'solusi': [
                'Tunggu halaman login dimuat sepenuhnya sebelum mengisi form.',
                'Refresh halaman browser dan coba login ulang.',
                'Jangan gunakan "Load Session Lama" jika session sudah lebih dari 24 jam.',
            ]
        },
        {
            'masalah': 'Scraping berhenti di tengah proses',
            'penyebab': 'Koneksi internet terputus atau session login expired.',
            'solusi': [
                'Periksa koneksi internet/LAN.',
                'Tutup browser, buka ulang, dan mulai scraping dari halaman terakhir.',
                'Jika session expired, login ulang terlebih dahulu.',
            ]
        },
        {
            'masalah': 'Data export Excel kosong atau corrupt',
            'penyebab': 'Tidak ada data yang berhasil di-scrape atau proses export gagal.',
            'solusi': [
                'Pastikan ada data di tabel hasil sebelum klik export.',
                'Coba export dalam format CSV sebagai alternatif.',
                'Refresh halaman dan coba export ulang.',
            ]
        },
        {
            'masalah': 'Aplikasi tidak bisa diakses dari komputer lain',
            'penyebab': 'Windows Firewall memblokir port atau IP belum dikonfigurasi.',
            'solusi': [
                'Jalankan Setup-NetworkAccess.bat sebagai Administrator.',
                'Pastikan file .env.local di folder frontend berisi IP server yang benar.',
                'Pastikan kedua komputer berada di jaringan yang sama.',
            ]
        },
        {
            'masalah': 'Pesan error "Module not found" / "Cannot find module"',
            'penyebab': 'Dependensi (library) belum terinstal.',
            'solusi': [
                'Untuk frontend: buka terminal, cd frontend, lalu npm install.',
                'Untuk SPP: cd spp-scraper, lalu pip install -r requirements.txt.',
                'Untuk DASTI: cd dasti-scraper, lalu pip install -r requirements.txt.',
                'Untuk SIPEDE: cd sipede-scraper/backend, lalu npm install.',
            ]
        },
    ]

    for item in problems:
        p = doc.add_paragraph()
        run = p.add_run(f'❌ Masalah: {item["masalah"]}')
        run.bold = True
        run.font.color.rgb = RGBColor(180, 0, 0)

        p = doc.add_paragraph()
        run = p.add_run('Penyebab: ')
        run.bold = True
        p.add_run(item['penyebab'])

        p = doc.add_paragraph()
        run = p.add_run('✅ Solusi:')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)
        for sol in item['solusi']:
            doc.add_paragraph(sol, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    doc.add_heading('8.3 FAQ (Pertanyaan yang Sering Diajukan)', level=2)

    faqs = [
        ('Apakah aplikasi ini menyimpan password login saya?',
         'Tidak. DASTA tidak menyimpan username atau password. '
         'Hanya cookies session yang disimpan untuk fitur "Load Session Lama" pada DASTI, '
         'dan cookies ini akan expired setelah 24 jam.'),
        ('Berapa banyak data yang bisa di-scrape dalam satu sesi?',
         'Tidak ada batasan dari sisi aplikasi. Jumlah data tergantung pada ketersediaan '
         'data di website sumber dan stabilitas koneksi jaringan. '
         'Anda dapat mengatur range halaman untuk membatasi jumlah data.'),
        ('Apakah saya bisa menggunakan aplikasi bersama orang lain?',
         'Ya, aplikasi dapat diakses dari beberapa komputer dalam jaringan yang sama. '
         'Namun, hanya satu proses scraping per modul yang dapat berjalan bersamaan.'),
        ('Data saya hilang setelah restart aplikasi, apa yang terjadi?',
         'Data scraping disimpan di database lokal dan seharusnya tidak hilang. '
         'Jika hilang, periksa file database di folder data/ masing-masing scraper. '
         'Data di Insight disimpan di localStorage browser — pastikan Anda tidak menghapus data browser.'),
        ('Apakah saya perlu koneksi internet untuk semua fitur?',
         'Tidak. SPDP hanya membutuhkan koneksi LAN. SIPEDE dan DASTI membutuhkan internet. '
         'Workspace dan Insight dapat digunakan offline dengan data yang sudah ada.'),
        ('Browser Chrome terbuka tapi halaman kosong, apa yang harus dilakukan?',
         'Hal ini biasanya karena versi ChromeDriver tidak sesuai dengan Chrome. '
         'Jalankan Fix-DastiChromeDriver.bat dan restart backend service.'),
        ('Bagaimana cara memperbarui aplikasi ke versi terbaru?',
         'Ambil versi terbaru dari repository, lalu jalankan ulang instalasi dependensi '
         '(npm install untuk frontend/SIPEDE, pip install -r requirements.txt untuk SPP/DASTI). '
         'Database dan data Anda tidak akan terpengaruh oleh pembaruan.'),
        ('Apa perbedaan antara Workspace dan Insight?',
         'Workspace adalah tempat untuk melihat, memfilter, dan mengekspor data mentah hasil scraping. '
         'Insight adalah tempat untuk menganalisis data surat dengan pengelompokan, '
         'visualisasi grafik, dan statistik lanjutan.'),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f'T: {q}')
        run.bold = True
        run.font.color.rgb = RGBColor(4, 78, 59)
        p = doc.add_paragraph()
        run = p.add_run('J: ')
        run.bold = True
        p.add_run(a)
        doc.add_paragraph()

    # ===================== FOOTER =====================
    doc.add_paragraph()
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run('— Akhir Dokumen —')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()
    closing2 = doc.add_paragraph()
    closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing2.add_run(
        f'Manual Book DASTA v2.0 — Panduan Penggunaan Aplikasi\n'
        f'Dibuat pada: {datetime.date.today().strftime("%d %B %Y")}\n'
        f'Klasifikasi: Internal'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # ===================== SAVE =====================
    output_path = 'Manual_Book_DASTA_Panduan_Penggunaan.docx'
    doc.save(output_path)
    print(f'✅ Manual book berhasil dibuat: {output_path}')
    return output_path


if __name__ == '__main__':
    create_manual()
