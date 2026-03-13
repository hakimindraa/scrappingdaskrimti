from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Title ---
title = doc.add_heading('Manual Book — Halaman Insight', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(30, 27, 75)

subtitle = doc.add_paragraph('Panduan Penggunaan Fitur Pengelompokan Data Surat')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph('')  # spacer

# ======================
# 1. Upload File Excel
# ======================
doc.add_heading('1. Upload File Excel', level=1)
doc.add_paragraph(
    'Fitur ini digunakan untuk meng-upload data surat masuk dan surat keluar dari file Excel (.xlsx / .xls).'
)

items = [
    'Klik tombol "Surat Masuk (.xlsx)" untuk upload data surat masuk.',
    'Klik tombol "Surat Keluar (.xlsx)" untuk upload data surat keluar.',
    'Format file harus berekstensi .xlsx atau .xls.',
    'Setelah upload berhasil, data akan tampil di tabel pengelompokan.',
    'Notifikasi akan muncul menampilkan jumlah data yang berhasil dibaca.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ======================
# 2. Kelompokkan Jenis & Asal Surat
# ======================
doc.add_heading('2. Kelompokkan Jenis & Asal Surat', level=1)
doc.add_paragraph(
    'Setelah upload, setiap jenis dan asal surat perlu dikelompokkan ke dalam kategori/kelompok yang telah ditentukan.'
)

doc.add_heading('Tab Jenis Surat', level=2)
items = [
    'Buka tab "Jenis Surat" pada tabel pengelompokan.',
    'Jenis yang sudah otomatis terkategori akan ditandai dengan centang hijau.',
    'Jenis yang belum terkategori akan menampilkan tombol "Assign".',
    'Klik "Assign" lalu pilih salah satu dari 7 kategori:',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

kategori = [
    'Surat Edaran',
    'Keputusan / Surat Perintah / Surat Tugas',
    'Undangan Internal / Eksternal',
    'Surat Pengantar / Lampiran',
    'Nota Dinas / Memorandum',
    'Laporan',
    'Biasa Internal / Eksternal',
]
for k in kategori:
    p = doc.add_paragraph(k, style='List Bullet 2')

doc.add_heading('Tab Asal Surat', level=2)
items = [
    'Buka tab "Asal Surat" pada tabel pengelompokan.',
    'Asal yang sudah otomatis terkelompok akan ditandai dengan centang hijau.',
    'Asal yang belum terkelompok akan menampilkan tombol "Assign".',
    'Klik "Assign" lalu pilih salah satu dari 8 kelompok:',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

kelompok = [
    'Kejaksaan',
    'Kemenkeu',
    'Pengadilan',
    'Perbankan',
    'Kepolisian dan BNN',
    'Gubernur / Pemda',
    'Instansi Lainnya / BUMN',
    'Aliansi Kemasyarakatan / Pribadi / LSM',
]
for k in kelompok:
    doc.add_paragraph(k, style='List Bullet 2')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Catatan: ')
run.bold = True
p.add_run('Mapping yang sudah di-assign akan tersimpan otomatis ke database. Saat upload file Excel baru, jenis/asal yang sama akan otomatis terkelompok tanpa perlu assign ulang.')

# ======================
# 3. Update Tahun & Bulan Per Baris
# ======================
doc.add_heading('3. Update Tahun & Bulan Per Baris', level=1)
doc.add_paragraph(
    'Fitur ini memungkinkan Anda mengubah bulan dan tahun untuk setiap baris data secara individual. '
    'Berguna jika data Excel memiliki tanggal yang tidak sesuai.'
)

doc.add_heading('Cara Menggunakan:', level=2)
items = [
    'Buka tab "Detail Data" di bagian bawah tabel pengelompokan.',
    'Pilih mode tampilan: "Per Kategori Jenis" atau "Per Kelompok Asal".',
    'Klik salah satu accordion kelompok (contoh: "Kejaksaan", "Surat Edaran") untuk melihat daftar baris.',
    'Pada kolom "Tanggal", terdapat dua dropdown: Bulan dan Tahun.',
    'Klik dropdown Bulan untuk mengubah bulan baris tersebut (Jan — Des).',
    'Klik dropdown Tahun untuk mengubah tahun baris tersebut.',
    'Perubahan langsung berlaku pada filter dan grafik tanpa perlu klik tombol simpan.',
    'Jika data disimpan ke database via tombol "Simpan ke Database", perubahan tahun/bulan juga ikut tersimpan.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Contoh penggunaan: ')
run.bold = True
p.add_run(
    'Surat masuk tercatat di bulan Januari, tetapi seharusnya masuk bulan Februari. '
    'Klik dropdown bulan pada baris tersebut, ubah dari "Jan" ke "Feb". '
    'Grafik dan tabel langsung ter-update.'
)

# ======================
# 4. Filter Bulan & Tahun
# ======================
doc.add_heading('4. Filter Bulan & Tahun', level=1)
doc.add_paragraph(
    'Anda dapat memfilter data berdasarkan rentang bulan dan tahun tertentu.'
)

items = [
    'Gunakan dropdown "Dari" dan "Sampai" di atas tabel untuk memfilter rentang bulan.',
    'Gunakan dropdown "Tahun" untuk memfilter tahun tertentu.',
    'Filter ini mempengaruhi semua tabel, grafik, dan statistik di bawahnya.',
    'Contoh: Filter Dari "Januari" Sampai "Maret" → hanya data bulan 1-3 yang tampil.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ======================
# 5. Simpan ke Database
# ======================
doc.add_heading('5. Simpan ke Database', level=1)
doc.add_paragraph(
    'Data surat yang sudah dikelompokkan dapat disimpan ke database SQLite untuk keperluan arsip.'
)

doc.add_heading('Cara Menyimpan:', level=2)
items = [
    'Pastikan semua data yang sedang difilter sudah terkelompok (tidak ada yang "Belum").',
    'Tombol "Simpan ke Database" akan berubah menjadi hijau (aktif) jika semua data sudah terkelompok.',
    'Klik tombol tersebut untuk menyimpan data + mapping ke database.',
    'Notifikasi hijau akan muncul jika berhasil.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Simpan Bertahap:', level=2)
items = [
    'Anda bisa menyimpan data bertahap per periode.',
    'Contoh: Filter Januari-Februari → kelompokkan semua → klik Simpan.',
    'Lalu filter Maret-April → kelompokkan semua → klik Simpan lagi.',
    'Data di database akan bertambah, bukan ditimpa.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Hapus Data dari Database:', level=2)
items = [
    'Klik tombol merah (ikon trash) di sebelah tombol Simpan.',
    'Konfirmasi penghapusan akan muncul.',
    'Data surat akan dihapus, tetapi mapping jenis→kategori dan asal→kelompok tetap tersimpan.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ======================
# 6. Aksi Per Baris
# ======================
doc.add_heading('6. Aksi Per Baris (Edit & Hapus)', level=1)
doc.add_paragraph(
    'Di tab Detail Data, setiap baris memiliki kolom "Aksi" dengan dua tombol:'
)

# Table for actions
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Ikon', 'Nama', 'Fungsi']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows_data = [
    ['Pensil', 'Edit', 'Mengubah kategori/kelompok baris tersebut secara individual (override per baris)'],
    ['Trash', 'Hapus', 'Menghapus baris dari kategori/kelompok saat ini (menjadikan "Belum")'],
]
for i, row_data in enumerate(rows_data):
    for j, value in enumerate(row_data):
        table.rows[i + 1].cells[j].text = value

# ======================
# 7. Catatan Penting
# ======================
doc.add_heading('7. Catatan Penting', level=1)

items = [
    'Mapping jenis→kategori dan asal→kelompok tersimpan otomatis setiap kali Anda assign. Upload Excel baru akan otomatis terkelompok.',
    'Data surat dari Excel hanya tersimpan jika Anda klik "Simpan ke Database".',
    'Setelah refresh halaman, tabel akan kosong. Upload ulang Excel untuk melihat data kembali (mapping tetap tersimpan).',
    'Database menggunakan SQLite (file dasti_data.db). Untuk melihat isi database, jalankan: python check-database.py',
    'Jika project diberikan ke orang lain, tidak perlu install database tambahan — SQLite sudah built-in di Python.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# --- Save ---
output_path = 'Manual_Book_Insight.docx'
doc.save(output_path)
print(f'Manual book berhasil dibuat: {output_path}')
